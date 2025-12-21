from PIL import Image, ImageDraw
from apps.api.app.extraction import extract_text_from_pdf, extract_text_from_image
import pytest
import docx
from pathlib import Path
from apps.api.app.extraction import extract_text_from_docx

def test_real_image_worker(tmp_path):
    image_path = tmp_path / "test.png"
    img = Image.new('RGB', (100, 30), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((10, 10), "Hello", fill=(0, 0, 0))
    img.save(image_path)
    result = extract_text_from_image(image_path)
    assert "Hello" in result

def test_real_pdf_worker():
    base_dir = Path(__file__).parent
    pdf_path = base_dir / "assets" / "Hello world.pdf"
    if not pdf_path.exists():
        pytest.skip("Skipping PDF test: hello.pdf not found in assets")
    result = extract_text_from_pdf(pdf_path)
    assert "Hello world!" in result.strip()

def test_real_docx_worker_success(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    file_path = tmp_path / "dynamic_test.docx"
    doc.save(file_path)

    result = extract_text_from_docx(file_path)

    assert "Hello World" in result
    assert "This is a test document." in result
    assert "\n" in result


def test_real_docx_worker_failure(tmp_path):
    bad_file = tmp_path / "corrupt.docx"
    bad_file.write_text("I am not a real zip file")

    with pytest.raises(ValueError) as excinfo:
        extract_text_from_docx(bad_file)

    assert "Error reading .docx file" in str(excinfo.value)